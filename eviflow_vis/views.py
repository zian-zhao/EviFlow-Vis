from django.shortcuts import render
from django.http import JsonResponse
import time
import base64
import binascii
import hashlib
import json
import requests
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple
from docx import Document
from pypdf import PdfReader
from django.views.decorators.csrf import csrf_exempt
import re
from .config import (
    DEEPSEEK_API_KEY,
    DEEPSEEK_API_URL,
    DEFAULT_CHART_THEME,
    CHART_ANIMATION_DURATION,
    CHART_ANIMATION_EASING,
    MAX_DATA_POINTS,
    MIN_DATA_POINTS,
    MAX_STRING_LENGTH
)
from .agents.master_agent import MasterAgent
from .chart_type_constants import normalize_chart_type_code
from . import strings_zh as SZH
from . import prompts_zh as PZH


def detect_text_language(text: str) -> str:
    """Simple zh/en detection for chart text rendering language."""
    text = text or ""
    cjk_count = len(re.findall(r'[\u4e00-\u9fff]', text))
    en_token_count = len(re.findall(r'\b[a-zA-Z]{2,}\b', text))
    if en_token_count >= max(6, cjk_count * 2):
        return 'en'
    return 'zh'


def contains_cjk(text: str) -> bool:
    return bool(re.search(r'[\u4e00-\u9fff]', text or ""))


def _translate_layout_text(text: str, language: str = 'zh') -> str:
    raw = str(text or "").strip()
    if language != 'en' or not raw:
        return raw
    if not contains_cjk(raw):
        return raw

    if raw in SZH.LAYOUT_ZH_TO_EN:
        return SZH.LAYOUT_ZH_TO_EN[raw]

    out = raw
    for zh, en in SZH.LAYOUT_ZH_FALLBACK_PAIRS:
        out = out.replace(zh, en)
    return out


def localize_layout_analysis(layout_analysis: Dict[str, Any], language: str = 'zh') -> Dict[str, Any]:
    data = json.loads(json.dumps(layout_analysis or {}, ensure_ascii=False))
    if language != 'en':
        return data

    for key in ('issues_detected', 'fixes_applied', 'suggestions'):
        vals = data.get(key, [])
        if isinstance(vals, list):
            data[key] = [_translate_layout_text(v, language='en') for v in vals]
    return data


NUMERIC_EVIDENCE_PATTERN = re.compile(
    r"(?:(?:\$|€|£|¥)\s*)?-?\d{1,3}(?:,\d{3})*(?:\.\d+)?(?:%|[kKmMbB]|万|亿)?|-?\d+(?:\.\d+)?%"
)


def _parse_numeric_token(token: str) -> float:
    t = (token or "").strip().lower().replace(",", "")
    scale = 1.0
    if t.endswith("%"):
        t = t[:-1]
    if t.endswith("k"):
        scale = 1_000.0
        t = t[:-1]
    elif t.endswith("m"):
        scale = 1_000_000.0
        t = t[:-1]
    elif t.endswith("b"):
        scale = 1_000_000_000.0
        t = t[:-1]
    elif t.endswith("万"):
        scale = 10_000.0
        t = t[:-1]
    elif t.endswith("亿"):
        scale = 100_000_000.0
        t = t[:-1]
    t = t.strip()
    if not t:
        raise ValueError("empty numeric token")
    return float(t) * scale


def _extract_numeric_evidence_values(text: str) -> List[float]:
    values: List[float] = []
    for m in NUMERIC_EVIDENCE_PATTERN.findall(text or ""):
        try:
            values.append(_parse_numeric_token(m))
        except Exception:
            continue
    return values


def _number_has_evidence(value: Any, evidence_values: List[float]) -> bool:
    if not isinstance(value, (int, float)):
        return True
    if not evidence_values:
        return False
    target = float(value)
    # Prefer exact / tight-tolerance matches against evidence-derived numbers.
    for ev in evidence_values:
        if abs(ev - target) <= max(1e-6, abs(target) * 1e-6):
            return True
    # Also accept rounded textual forms (integers / half-up rounding).
    for ev in evidence_values:
        if round(ev, 2) == round(target, 2) or round(ev) == round(target):
            return True
    return False


def _collect_series_numeric_points(series_item: Dict[str, Any]) -> int:
    total = 0
    data = series_item.get("data", [])
    if not isinstance(data, list):
        return 0
    for point in data:
        if isinstance(point, (int, float)):
            total += 1
        elif isinstance(point, dict):
            v = point.get("value")
            if isinstance(v, (int, float)):
                total += 1
            elif isinstance(v, list):
                total += sum(1 for x in v if isinstance(x, (int, float)))
        elif isinstance(point, list):
            total += sum(1 for x in point if isinstance(x, (int, float)))
    return total


def sanitize_chart_config_by_evidence(chart_config: Dict[str, Any], description: str) -> Tuple[Dict[str, Any], Dict[str, Any]]:
    """
    Drop unsupported numeric points that cannot be traced to source text.
    Returns sanitized config + evidence report.
    """
    cfg = json.loads(json.dumps(chart_config or {}, ensure_ascii=False))
    evidence_values = _extract_numeric_evidence_values(description)
    report = {
        "source_numeric_count": len(evidence_values),
        "total_numeric_points": 0,
        "kept_numeric_points": 0,
        "dropped_numeric_points": 0,
        "drop_ratio": 0.0,
    }

    series_list = cfg.get("series", [])
    if isinstance(series_list, dict):
        series_list = [series_list]
    if not isinstance(series_list, list):
        return cfg, report

    sanitized_series = []
    for s in series_list:
        if not isinstance(s, dict):
            continue
        report["total_numeric_points"] += _collect_series_numeric_points(s)
        data = s.get("data", [])
        if not isinstance(data, list):
            sanitized_series.append(s)
            continue

        new_data = []
        for point in data:
            keep_point = True
            kept_in_point = 0
            dropped_in_point = 0

            if isinstance(point, (int, float)):
                keep_point = _number_has_evidence(point, evidence_values)
                kept_in_point = 1 if keep_point else 0
                dropped_in_point = 0 if keep_point else 1
            elif isinstance(point, dict):
                v = point.get("value")
                if isinstance(v, (int, float)):
                    keep_point = _number_has_evidence(v, evidence_values)
                    kept_in_point = 1 if keep_point else 0
                    dropped_in_point = 0 if keep_point else 1
                elif isinstance(v, list):
                    filtered = []
                    for x in v:
                        if isinstance(x, (int, float)):
                            if _number_has_evidence(x, evidence_values):
                                filtered.append(x)
                                kept_in_point += 1
                            else:
                                dropped_in_point += 1
                        else:
                            filtered.append(x)
                    if not filtered and any(isinstance(x, (int, float)) for x in v):
                        keep_point = False
                    else:
                        point = {**point, "value": filtered}
                else:
                    keep_point = True
            elif isinstance(point, list):
                filtered = []
                for x in point:
                    if isinstance(x, (int, float)):
                        if _number_has_evidence(x, evidence_values):
                            filtered.append(x)
                            kept_in_point += 1
                        else:
                            dropped_in_point += 1
                    else:
                        filtered.append(x)
                if not filtered and any(isinstance(x, (int, float)) for x in point):
                    keep_point = False
                else:
                    point = filtered
            else:
                keep_point = True

            report["kept_numeric_points"] += kept_in_point
            report["dropped_numeric_points"] += dropped_in_point
            if keep_point:
                new_data.append(point)

        sanitized = {**s, "data": new_data}
        if new_data:
            sanitized_series.append(sanitized)

    cfg["series"] = sanitized_series
    total = report["total_numeric_points"]
    dropped = report["dropped_numeric_points"]
    report["drop_ratio"] = (dropped / total) if total > 0 else 0.0
    return cfg, report


def infer_chart_type_scores(description: str) -> Dict[str, float]:
    """Heuristic chart-type scoring to stabilize type selection."""
    text = (description or "")
    lower = text.lower()
    numeric_count = len(NUMERIC_EVIDENCE_PATTERN.findall(text))
    has_time = bool(
        re.search(r"\b(q[1-4]|week\s*\d+|month|year|daily|weekly|monthly|quarter|timeline)\b", lower)
        or re.search(SZH.HEURISTIC_ZH_TIME, text)
    )
    has_compare = bool(
        re.search(r"\b(compare|versus|rank|higher|lower|top|bottom)\b", lower)
        or re.search(SZH.HEURISTIC_ZH_COMPARE, text)
    )
    has_share = bool(
        re.search(r"\b(share|ratio|proportion|percentage|percent)\b", lower)
        or re.search(SZH.HEURISTIC_ZH_SHARE, text)
    )
    has_relation = bool(
        re.search(r"\b(correlation|relationship|association)\b", lower)
        or re.search(SZH.HEURISTIC_ZH_RELATION, text)
    )

    scores = {
        "line": 0.0,
        "bar": 0.0,
        "scatter": 0.0,
        "pie": 0.0,
        "radar": 0.0,
        "funnel": 0.0,
        "tree": 0.0,
        "gantt": 0.0,
        "combo": 0.0,
    }
    if has_time:
        scores["line"] += 30
        scores["bar"] += 10
    if has_compare:
        scores["bar"] += 28
        scores["line"] += 12
    if has_share:
        scores["pie"] += 26
        scores["bar"] += 12
    if has_relation:
        scores["scatter"] += 24
    if numeric_count >= 3:
        scores["bar"] += 14
        scores["line"] += 14
    elif numeric_count <= 1:
        scores["tree"] += 8
        scores["radar"] += 6
    if has_time and has_compare:
        scores["combo"] += 26
    if numeric_count >= 4 and (has_time or has_compare):
        scores["combo"] += 10

    # Give small priors to robust charts.
    scores["bar"] += 4
    scores["line"] += 4
    return scores


def choose_chart_type(description: str, llm_type: str) -> Tuple[str, Dict[str, float]]:
    """Blend LLM-selected type with heuristic scores for better reliability."""
    scores = infer_chart_type_scores(description)
    llm_code = normalize_chart_type_code(llm_type or "")
    if llm_code in scores:
        scores[llm_code] += 12
    ranked = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
    top = ranked[0][0] if ranked else (llm_code or "bar")
    return (top, dict(ranked))


def enforce_chart_type_consistency(chart_config: Dict[str, Any], chart_type: str) -> Dict[str, Any]:
    """Constrain output option to selected chart type to reduce mismatch."""
    cfg = json.loads(json.dumps(chart_config or {}, ensure_ascii=False))
    selected = (chart_type or "bar").lower()
    series = cfg.get("series")
    if isinstance(series, dict):
        series = [series]
    if not isinstance(series, list):
        return cfg

    # Composite (combo): keep per-series types (line/bar/scatter mix); do not flatten.
    if selected == "combo":
        for s in series:
            if isinstance(s, dict) and not s.get("type"):
                s["type"] = "bar"
        cfg["series"] = series
        if not cfg.get("xAxis"):
            cfg.setdefault("xAxis", {"type": "category"})
        if cfg.get("yAxis") is None:
            cfg.setdefault("yAxis", {"type": "value"})
        return cfg

    # For cartesian charts, make series types consistent.
    if selected in {"line", "bar", "scatter"}:
        for s in series:
            if isinstance(s, dict):
                s["type"] = selected
        cfg["series"] = series
        cfg.setdefault("xAxis", {"type": "category"})
        cfg.setdefault("yAxis", {"type": "value"})
        return cfg

    # For pie / radar / funnel / tree / gantt, force first series type.
    if series and isinstance(series[0], dict):
        series[0]["type"] = selected
    cfg["series"] = series
    return cfg

def analyze_with_llm(description):
    """Call DeepSeek chat API to infer chart skeleton from a natural-language brief."""
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    
    prompt = PZH.analyze_data_description_user_zh(
        description,
        DEFAULT_CHART_THEME,
        CHART_ANIMATION_DURATION,
        CHART_ANIMATION_EASING,
        MIN_DATA_POINTS,
        MAX_DATA_POINTS,
        MAX_STRING_LENGTH,
    )

    data = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": PZH.analyze_data_description_system_zh()},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.7,
        "max_tokens": 2000
    }

    try:
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data)
        response.raise_for_status()
        result = response.json()
        
        llm_response = result['choices'][0]['message']['content']
        json_start = llm_response.find('{')
        json_end = llm_response.rfind('}') + 1
        if json_start >= 0 and json_end > json_start:
            json_str = llm_response[json_start:json_end]
            return json.loads(json_str)
        else:
            raise ValueError(SZH.ERR_CANNOT_EXTRACT_JSON_FROM_LLM)
            
    except Exception as e:
        print(f"analyze_with_llm failed: {e}")
        return None


@csrf_exempt
def chart_workbench(request):
    if request.method == 'POST':
        try:
            data: Dict[str, Any] = {}
            # Parse chart request payload (JSON or form).
            if request.content_type == 'application/json':
                data = json.loads(request.body)
                description = data.get('data', '')
            else:
                description = request.POST.get('description', '')

            if not description:
                return JsonResponse({'error': SZH.ERR_PROVIDE_DATA_DESCRIPTION})

            language = detect_text_language(description)
            preferred_raw = (data.get("preferred_chart_type") or data.get("chart_type") or "").strip()
            preferred_code = normalize_chart_type_code(preferred_raw) if preferred_raw else ""
            master_agent = MasterAgent()
            orchestration_result = master_agent.orchestrate_graph_generation(
                description=description,
                language=language,
                generate_chart_config_fn=generate_chart_config,
                enforce_chart_type_consistency_fn=enforce_chart_type_consistency,
                sanitize_chart_config_by_evidence_fn=sanitize_chart_config_by_evidence,
                chart_type_override=preferred_code or None,
            )
            if orchestration_result.get('error'):
                return JsonResponse({'error': orchestration_result.get('error')})

            chart_config = orchestration_result['chart_config']
            selected_chart_type = orchestration_result['chart_type']
            layout_analysis = localize_layout_analysis(orchestration_result['layout_analysis'], language=language)
            evidence_report = orchestration_result['evidence_report']
            type_scores = orchestration_result['chart_type_scores']
            analysis_result = orchestration_result['analysis']
            chart_type_recommendations = build_chart_type_recommendations_from_scores(
                type_scores,
                selected_chart_type,
                description,
                language,
            )
            analysis_summary = build_generation_aligned_analysis_summary(language)

            # Build ECharts init snippet for the browser.
            script = f"""
            var myChart = echarts.init(document.getElementById('chart'));
            var option = {json.dumps(chart_config, ensure_ascii=False)};
            myChart.setOption(option);
            """

            return JsonResponse({
                'script': script,
                'chart_type': selected_chart_type,
                'title': chart_config.get('title', {}).get('text', 'Data Visualization' if language == 'en' else SZH.UI_DEFAULT_CHART_TITLE_ZH),
                'description': chart_config.get('title', {}).get('subtext', ''),
                'analysis': analysis_result,
                'layout_analysis': layout_analysis,  # overlap / repair hints for the UI
                'language': language,
                'evidence_report': evidence_report,
                'chart_type_scores': type_scores,
                'chart_type_recommendations': chart_type_recommendations,
                'analysis_summary': analysis_summary,
            })

        except json.JSONDecodeError:
            return JsonResponse({'error': SZH.ERR_INVALID_JSON})
        except Exception as e:
            return JsonResponse({'error': str(e)})

    return render(request, 'viz_lab/chart_workbench.html')


def chart_layout_report(request):
    """Standalone page: auto-layout charts with source text, manual tweaks, export."""
    return render(request, 'viz_lab/chart_layout_report.html')


def _docx_add_multiline_paragraph(doc, text, italic=False):
    """Insert a paragraph preserving newline characters as soft line breaks."""
    from docx.enum.text import WD_BREAK

    if text is None:
        text = ''
    lines = text.split('\n')
    para = doc.add_paragraph()
    if lines:
        r0 = para.add_run(lines[0])
        r0.italic = italic
        for line in lines[1:]:
            para.add_run().add_break(WD_BREAK.LINE)
            rr = para.add_run(line)
            rr.italic = italic


@csrf_exempt
def chart_layout_export_docx(request):
    """
    Build a .docx whose body follows mergedParts (full text flow + figures).
    Expects JSON: { "mergedParts": [...], "chartImages": { "0": "<base64 or data URL>", ... } }
    """
    from django.http import HttpResponse
    from docx.shared import Inches, Pt

    if request.method != 'POST':
        return JsonResponse({'error': SZH.ERR_POST_ONLY}, status=405)

    try:
        data = json.loads(request.body.decode('utf-8'))
    except (json.JSONDecodeError, UnicodeDecodeError):
        return JsonResponse({'error': SZH.ERR_INVALID_JSON_BODY}, status=400)

    merged_parts = data.get('mergedParts')
    chart_images = data.get('chartImages') or {}
    if not isinstance(merged_parts, list):
        return JsonResponse({'error': SZH.ERR_MERGED_PARTS_NOT_ARRAY}, status=400)

    doc = Document()
    figure_n = 0

    for part in merged_parts:
        if not isinstance(part, dict):
            continue
        ptype = part.get('type')
        if ptype == 'plain':
            _docx_add_multiline_paragraph(doc, part.get('text') or '', italic=False)
        elif ptype == 'excerpt':
            _docx_add_multiline_paragraph(doc, part.get('text') or '', italic=True)
        elif ptype == 'orphan_note':
            p = doc.add_paragraph(part.get('text') or '')
            for r in p.runs:
                r.italic = True
                r.font.size = Pt(9)
        elif ptype == 'chart':
            ci = part.get('chartIndex')
            key = str(ci) if ci is not None else ''
            raw = chart_images.get(key)
            if raw is None and ci is not None:
                raw = chart_images.get(ci)
            figure_n += 1
            if raw:
                if isinstance(raw, str) and 'base64,' in raw:
                    raw = raw.split('base64,', 1)[-1]
                elif isinstance(raw, str) and raw.startswith('data:'):
                    raw = raw.split(',', 1)[-1]
                try:
                    blob = base64.b64decode(raw, validate=False)
                except (binascii.Error, ValueError):
                    blob = None
                if blob:
                    try:
                        doc.add_picture(BytesIO(blob), width=Inches(5.9))
                    except Exception:
                        p = doc.add_paragraph('(Image decode failed; skipped.)')
                        for r in p.runs:
                            r.italic = True
            cap_text = (part.get('caption') or '').strip()
            if not cap_text:
                cap_text = f'Fig. {figure_n}'
            cap = doc.add_paragraph(cap_text)
            for r in cap.runs:
                r.italic = True
                r.font.size = Pt(9)
        else:
            continue

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    resp = HttpResponse(
        bio.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    resp['Content-Disposition'] = 'attachment; filename="document-with-figures.docx"'
    return resp


@csrf_exempt
def extract_chart_text_segments(request):
    """Accept uploaded text and return chartable spans (txt/md/csv/docx/pdf)."""
    if request.method != 'POST':
        return JsonResponse({'error': SZH.ERR_POST_ONLY_ALT})

    upload = request.FILES.get('text_file')
    if not upload:
        return JsonResponse({'error': SZH.ERR_UPLOAD_TEXT_FILE})

    filename = upload.name.lower()
    if not (filename.endswith('.txt') or filename.endswith('.md') or filename.endswith('.csv') or filename.endswith('.docx') or filename.endswith('.pdf')):
        return JsonResponse({'error': SZH.ERR_UNSUPPORTED_TEXT_EXT})

    try:
        profile = (request.POST.get('confidence_profile') or 'balanced').strip().lower()
        if profile not in ('strict', 'balanced', 'loose'):
            profile = 'balanced'

        raw_bytes = upload.read()
        if filename.endswith('.docx'):
            doc = Document(BytesIO(raw_bytes))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            content = '\n'.join(paragraphs)
        elif filename.endswith('.pdf'):
            reader = PdfReader(BytesIO(raw_bytes))
            page_texts = []
            for page in reader.pages:
                text = page.extract_text() or ''
                text = text.strip()
                if text:
                    page_texts.append(text)
            content = '\n'.join(page_texts)
        else:
            content = raw_bytes.decode('utf-8', errors='ignore')

        if not content.strip():
            return JsonResponse({'error': SZH.ERR_EMPTY_FILE})

        master_agent = MasterAgent()
        result = master_agent.orchestrate_segment_extraction(content, profile=profile)
        segments = result.get('segments', [])
        return JsonResponse({
            'text': content,
            'segments': segments,
            'profile': profile
        })
    except Exception as e:
        return JsonResponse({'error': SZH.err_text_parse_failed(e)})

@csrf_exempt
def analyze_chart_types(request):
    """JSON endpoint: LLM-assisted chart-type ranking for a single brief."""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            description = data.get('description', '')
            backend_analysis = data.get('backend_analysis', {})
            language = data.get('language') or detect_text_language(description)

            if not description:
                return JsonResponse({'error': SZH.ERR_PROVIDE_DATA_DESCRIPTION})

            # LLM-assisted chart-type ranking
            analysis_result = analyze_chart_types_with_llm(description, backend_analysis, language=language)
            
            if analysis_result:
                return JsonResponse(analysis_result)
            else:
                return JsonResponse({'error': SZH.ERR_LLM_ANALYSIS_FAILED})

        except json.JSONDecodeError:
            return JsonResponse({'error': SZH.ERR_INVALID_JSON})
        except Exception as e:
            return JsonResponse({'error': str(e)})

    return JsonResponse({'error': SZH.ERR_POST_ONLY_ALT})


@csrf_exempt
def check_chart_layout_overlap(request):
    """POST JSON with chart_config (+ optional description) to run overlap QC."""
    if request.method != 'POST':
        return JsonResponse({'error': SZH.ERR_POST_ONLY_ALT})

    try:
        data = json.loads(request.body)
        chart_config = data.get('chart_config') or {}
        description = data.get('description', '')
        auto_fix = bool(data.get('auto_fix', False))
        language = data.get('language') or detect_text_language(description)

        if not isinstance(chart_config, dict) or not chart_config:
            return JsonResponse({'error': SZH.ERR_CHART_CONFIG_REQUIRED})

        from eviflow_vis.agents.quality_control_agent import QualityControlAgent

        qc_agent = QualityControlAgent()
        layout_result = qc_agent.check_layout_overlaps_sync(chart_config, description)

        layout_analysis = {
            'issues_detected': layout_result.get('issues', []),
            'fixes_applied': [],
            'overlap_score': layout_result.get('overlap_score', 0),
            'suggestions': layout_result.get('suggestions', []),
        }

        fixed_config = None
        if auto_fix and layout_result.get('has_overlaps', False):
            fix_result = qc_agent.auto_fix_overlaps_sync(chart_config, layout_result.get('issues', []))
            fixed_config = fix_result.get('fixed_config')
            layout_analysis['fixes_applied'] = fix_result.get('fixes_applied', [])

        layout_analysis = localize_layout_analysis(layout_analysis, language=language)

        resp = {
            'layout_analysis': layout_analysis,
            'has_overlaps': bool(layout_result.get('has_overlaps', False)),
        }
        if fixed_config is not None:
            resp['fixed_config'] = fixed_config
        return JsonResponse(resp)
    except json.JSONDecodeError:
        return JsonResponse({'error': SZH.ERR_INVALID_JSON})
    except Exception as e:
        return JsonResponse({'error': str(e)})

def generate_chart_config(description, chart_type, analysis_result, language='zh'):
    """Build an ECharts option JSON via the configured chat model."""
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    is_combo = str(chart_type or "").strip().lower() == "combo"
    combo_rules_en = """
13. COMPOSITE CHART (combo): You design the full option. Different series may use different series.type among line, bar, scatter, and area on a shared category or time axis.
14. When there are two or more distinct numeric metrics on the same categories or years, use at least TWO different series.type values (e.g. bar + line, or bar + scatter). Strong default: bar for the larger-scale metric and line for the smaller-scale or rate-like metric; avoid using only identical types (e.g. two plain areas) unless the source text explicitly asks for that style.
15. You may set yAxis to an array and assign yAxisIndex per series for dual-axis when two metrics have very different scales.
16. Do not mix pie, radar, funnel, tree, or gantt series with cartesian series in the same chart.
17. Give each series a clear name for the legend; align colors for readability.
18. Use available data fully; do not discard valid numeric points from source text.
"""
    single_type_rule_en = f'13. IMPORTANT: every series[].type MUST be "{chart_type}" unless the data structure makes that impossible.'
    single_type_rule_zh = SZH.single_series_type_rule_zh(chart_type)

    if language == 'en':
        prompt = f"""Generate a complete ECharts option in valid JSON.

Data description: {description}
Chart type: {chart_type}
Analysis result: {json.dumps(analysis_result, ensure_ascii=False)}

Requirements:
1. Include title and subtitle.
2. Include axes when needed.
3. Include series, legend, and tooltip.
4. Use theme: {DEFAULT_CHART_THEME}
5. Animation duration: {CHART_ANIMATION_DURATION}ms
6. Animation easing: {CHART_ANIMATION_EASING}
7. Data points limit: {MIN_DATA_POINTS} to {MAX_DATA_POINTS}
8. String length limit: {MAX_STRING_LENGTH}
9. All human-readable labels (title, legend text, axis labels) should be in English.
10. ZERO-FABRICATION POLICY: never invent or extrapolate numeric values not explicitly present in source text.
11. If a value is missing in source text, omit that datapoint (or set null), do not guess.
12. Keep only evidence-backed values from source description.
{(combo_rules_en if is_combo else single_type_rule_en + chr(10) + "14. Use available data fully; do not discard valid numeric points from source text.")}

Return JSON only."""
        system_prompt = "You are an expert in data visualization and ECharts option design. Output JSON only."
        fix_system_prompt = "You are an expert in data visualization. Output valid JSON only."
    else:
        prompt = PZH.generate_chart_config_user_zh(
            description,
            chart_type,
            analysis_result,
            DEFAULT_CHART_THEME,
            CHART_ANIMATION_DURATION,
            CHART_ANIMATION_EASING,
            MIN_DATA_POINTS,
            MAX_DATA_POINTS,
            MAX_STRING_LENGTH,
            is_combo,
            SZH.COMBO_RULES_ZH,
            single_type_rule_zh,
        )
        system_prompt = PZH.generate_chart_config_system_zh()
        fix_system_prompt = PZH.generate_chart_config_fix_system_zh()

    data = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.7,
        "max_tokens": 2000
    }

    from eviflow_vis.agents.quality_control_agent import QualityControlAgent
    qc_agent = QualityControlAgent()

    def call_llm_func(fix_prompt):
        """Inner helper for QC self-correction rounds."""
        fix_data = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "system", "content": fix_system_prompt},
                {"role": "user", "content": fix_prompt}
            ],
            "temperature": 0.3,
            "max_tokens": 1500
        }
        try:
            fix_response = requests.post(DEEPSEEK_API_URL, headers=headers, json=fix_data)
            fix_response.raise_for_status()
            fix_result = fix_response.json()
            return fix_result['choices'][0]['message']['content']
        except Exception as e:
            print(f"chart_config self-fix LLM call failed: {e}")
            return ""

    try:
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data)
        response.raise_for_status()
        result = response.json()
        llm_response = result['choices'][0]['message']['content']
        chart_config = qc_agent.parse_llm_chart_config(llm_response, call_llm_func=call_llm_func)
        if chart_config and not chart_config.get('error'):
            return chart_config
        else:
            print("raw LLM chart response:", llm_response)
            print("parse/fix error:", chart_config.get('error'))
            return None
    except Exception as e:
        print(f"generate_chart_config failed: {e}")
        return None

def analyze_chart_types_with_llm(description, backend_analysis, language='zh'):
    """Rank chart-type candidates with the chat model (confidence must sum to 100%)."""
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    
    if language == 'en':
        prompt = f"""Analyze the data description and recommend the most suitable chart types.

Description: {description}
Backend analysis: {json.dumps(backend_analysis, ensure_ascii=False)}

Return JSON:
{{
    "recommendations": [
        {{
            "type": "chart type code",
            "name": "English chart type name",
            "confidence": 0-100,
            "reason": "professional reason in English",
            "matched_keywords": ["keyword1", "keyword2"],
            "suitability_score": 0-100
        }}
    ],
    "analysis_summary": {{
        "data_characteristics": "brief English summary",
        "visualization_goals": "brief English summary",
        "recommended_approach": "brief English summary"
    }}
}}

Constraints:
1. Confidence must reflect real suitability.
2. Total confidence must sum to 100%.
3. Sort recommendations by confidence descending.
4. Return valid JSON only.
5. Allowed type codes include line, bar, scatter, pie, radar, funnel, tree, gantt, area, heatmap, bubble, candlestick, and combo (composite chart: multiple series with types among line/bar/scatter, dual y-axis allowed when needed)."""
        system_prompt = "You are a professional data visualization expert. Ensure confidence sums to 100 and avoid fabricated scores."
    else:
        prompt = PZH.chart_type_recommendation_user_zh(description, backend_analysis)
        system_prompt = PZH.chart_type_recommendation_system_zh()

    data = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.3,
        "max_tokens": 2000
    }

    try:
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data)
        response.raise_for_status()
        result = response.json()
        
        llm_response = result['choices'][0]['message']['content']
        json_start = llm_response.find('{')
        json_end = llm_response.rfind('}') + 1
        if json_start >= 0 and json_end > json_start:
            analysis_result = json.loads(llm_response[json_start:json_end])
            
            # Validate / normalize recommendation payload
            if 'recommendations' in analysis_result and isinstance(analysis_result['recommendations'], list):
                recommendations = analysis_result['recommendations']
                
                # 确保置信度在合理范围内并归一化
                total_confidence = 0
                valid_recommendations = []
                
                for rec in recommendations:
                    if not isinstance(rec, dict):
                        continue
                    if not rec.get('type'):
                        continue
                    if 'confidence' in rec and isinstance(rec['confidence'], (int, float)):
                        confidence = max(0, min(100, float(rec['confidence'])))
                        rec['confidence'] = confidence
                        total_confidence += confidence
                    elif isinstance(rec.get('suitability_score'), (int, float)):
                        confidence = max(0, min(100, float(rec['suitability_score'])))
                        rec['confidence'] = confidence
                        total_confidence += confidence
                    valid_recommendations.append(rec)
                
                # 如果总置信度不为100%，进行归一化
                if valid_recommendations and total_confidence > 0:
                    for rec in valid_recommendations:
                        rec['confidence'] = round((rec['confidence'] / total_confidence) * 100, 1)
                elif valid_recommendations:
                    # 没有可用置信度时，使用后端启发式分布，避免机械均分
                    fallback_scores = infer_chart_type_scores(description)
                    scored = []
                    for rec in valid_recommendations:
                        chart_type = str(rec.get('type', '')).lower()
                        # rank 偏置用于打破平分；值很小，仅在分数近似时生效
                        base_score = float(fallback_scores.get(chart_type, 0.0))
                        scored.append((rec, base_score))
                    scored.sort(key=lambda x: x[1], reverse=True)
                    score_sum = sum(max(0.0, s) for _, s in scored)
                    if score_sum <= 0:
                        # 最后兜底：给递减分布，而不是25/25/25/25
                        descending = [40.0, 30.0, 20.0, 10.0]
                        for i, (rec, _) in enumerate(scored):
                            rec['confidence'] = descending[i] if i < len(descending) else 0.0
                    else:
                        for rec, score in scored:
                            rec['confidence'] = round((max(0.0, score) / score_sum) * 100, 1)
                
                # 确保其他必要字段存在
                for rec in valid_recommendations:
                    if not rec.get('type'):
                        rec['type'] = 'bar'
                    if not rec.get('name'):
                        rec['name'] = get_chart_type_name(rec['type'], language=language)
                    elif language == 'en' and contains_cjk(str(rec.get('name', ''))):
                        rec['name'] = get_chart_type_name(rec['type'], language=language)
                    matched = rec.get('matched_keywords') or []
                    if not matched:
                        matched = _infer_matched_keywords(description, rec['type'], language=language)
                    rec['matched_keywords'] = matched[:4]

                    reason_text = str(rec.get('reason') or '').strip()
                    if (
                        not reason_text
                        or "common practice" in reason_text.lower()
                        or "best fit" in reason_text.lower()
                        or SZH.SPURIOUS_REASON_PHRASE_ZH in reason_text
                    ):
                        rec['reason'] = _build_recommendation_reason(
                            description=description,
                            chart_type=rec['type'],
                            confidence=float(rec.get('confidence', 0)),
                            matched_keywords=rec['matched_keywords'],
                            language=language,
                        )
                    if 'suitability_score' not in rec:
                        rec['suitability_score'] = rec['confidence']
                
                # 按置信度排序
                valid_recommendations.sort(key=lambda x: x.get('confidence', 0), reverse=True)
                
                # 限制推荐数量为4个
                analysis_result['recommendations'] = valid_recommendations[:4]

                # 当置信度分布过平时，按启发式再次拉开差异，避免4个都25%
                if analysis_result['recommendations']:
                    conf_values = [float(rec.get('confidence', 0)) for rec in analysis_result['recommendations']]
                    if max(conf_values) - min(conf_values) < 1.0:
                        fallback_scores = infer_chart_type_scores(description)
                        rescored = []
                        for idx, rec in enumerate(analysis_result['recommendations']):
                            chart_type = str(rec.get('type', '')).lower()
                            tie_break = max(0.0, 1.0 - idx * 0.05)
                            score = max(0.0, float(fallback_scores.get(chart_type, 0.0))) + tie_break
                            rescored.append((rec, score))
                        score_sum = sum(score for _, score in rescored)
                        if score_sum > 0:
                            for rec, score in rescored:
                                rec['confidence'] = round((score / score_sum) * 100, 1)
                
                # 最终验证总和是否为100%
                final_total = sum(rec.get('confidence', 0) for rec in analysis_result['recommendations'])
                if abs(final_total - 100) > 0.1:  # 允许0.1的误差
                    # 重新归一化
                    for rec in analysis_result['recommendations']:
                        rec['confidence'] = round((rec['confidence'] / final_total) * 100, 1) if final_total > 0 else 0.0
                
                return analysis_result
            else:
                raise ValueError(SZH.ERR_LLM_RESULT_SHAPE_INVALID)
        else:
            raise ValueError(SZH.ERR_CANNOT_EXTRACT_JSON_FROM_LLM)
            
    except Exception as e:
        print(f"analyze_chart_types_with_llm failed: {e}")
        return None

def get_chart_type_name(chart_type, language='zh'):
    """Resolve a stable display name for a chart-type code."""
    table = SZH.CHART_TYPE_NAMES_EN if language == 'en' else SZH.CHART_TYPE_NAMES_ZH
    return table.get(chart_type, chart_type)


def _infer_matched_keywords(description: str, chart_type: str, language: str = 'zh') -> List[str]:
    text = description or ""
    lower = text.lower()
    keyword_map = SZH.KEYWORD_MAP_BY_CHART_TYPE
    candidates = keyword_map.get(str(chart_type).lower(), [])
    hits = []
    for kw in candidates:
        if kw.isascii():
            if kw in lower:
                hits.append(kw)
        else:
            if kw in text:
                hits.append(kw)
    if hits:
        return hits[:4]
    return (SZH.DEFAULT_MATCHED_KEYWORDS_EN if language == 'en' else SZH.DEFAULT_MATCHED_KEYWORDS_ZH)


def _build_recommendation_reason(
    description: str,
    chart_type: str,
    confidence: float,
    matched_keywords: List[str],
    language: str = 'zh',
) -> str:
    chart_name = get_chart_type_name(chart_type, language=language)
    kw = ", ".join((matched_keywords or [])[:3]) if matched_keywords else (
        "semantic cues" if language == 'en' else PZH.semantic_fallback_label_zh()
    )
    numeric_count = len(NUMERIC_EVIDENCE_PATTERN.findall(description or ""))
    if language == 'en':
        return (
            f"{chart_name} is preferred because cues [{kw}] align with this encoding; "
            f"detected numeric evidence count={numeric_count}; confidence={round(confidence, 1)}%."
        )
    return SZH.fmt_recommendation_reason_zh(chart_name, kw, numeric_count, confidence)


_RUNNER_OPENING_EN = (
    "If you want a different lens, ",
    "Still defensible: ",
    "Another credible read: ",
    "As a runner-up framing, ",
)

# Short pools; combined with rank + MD5(seed) they stay varied when LLM is skipped.
_REASON_EN: Dict[str, Tuple[str, ...]] = {
    "line": (
        "Ordered readings here behave like a path—lines foreground slope, reversals, and which series crosses which.",
        "When the question is “how did it move across steps or periods?”, a line keeps continuity without implying unrelated buckets.",
        "Several aligned numeric ticks invite a strip-chart read: shape first, clutter second.",
    ),
    "bar": (
        "Detached categories with one main number per label read fastest as bars; the eye compares lengths without reading slopes.",
        "If the payoff is “who is tallest?” across a modest set of buckets, equal-width columns keep pairwise judgment honest.",
        "Nominal labels plus magnitudes are the classic bar story—clear baselines and room for readable tick labels.",
    ),
    "scatter": (
        "When two numeric channels might co-vary, scatter encodings expose clusters, outliers, and gentle trends at once.",
        "Pairwise measurements without a strict time axis often deserve a cloud of points rather than connected segments.",
        "If you care about joint behavior more than rank order, plotting x-y pairs avoids implying a false sequence.",
    ),
    "pie": (
        "Share-of-whole language in the excerpt fits part-to-whole encoding; pies (or donuts) make relative weight visceral.",
        "When components must sum to a meaningful total, angular slices emphasize composition rather than precise rank gaps.",
        "A handful of parts that should read as percentages map naturally to wedge metaphors for non-technical audiences.",
    ),
    "radar": (
        "Several traits on the same entities call for a spider plot—one glance shows skewed strengths and weak dimensions.",
        "Multi-axis profile comparisons (scores, rubrics) stay compact on a radar instead of many separate tiny charts.",
        "When every axis shares a comparable scale, folding them into one polygon highlights imbalance patterns.",
    ),
    "funnel": (
        "Sequential stages with attrition read as a funnel: width encodes surviving volume at each gate.",
        "Pipeline metaphors with shrinking counts between steps are easier to scan top-to-bottom than as disconnected bars.",
        "If the story is drop-off through ordered phases, tapering bands match how people narrate conversions.",
    ),
    "tree": (
        "Nested labels or parent/child rollups in the text map cleanly to a tree—structure beats magnitude here.",
        "When hierarchy matters more than exact leaf values, branching layout preserves containment at a glance.",
        "If categories nest inside categories, a treemap or tree view prevents misreading unrelated siblings as ordered time.",
    ),
    "gantt": (
        "Spans with start/stop flavor across a calendar suggest a Gantt-style lane chart for overlap and slack.",
        "When tasks compete for the same timeline, horizontal bars on a date axis expose concurrency and bottlenecks.",
        "Milestone-heavy narratives benefit from explicit duration bars instead of isolated point markers.",
    ),
    "combo": (
        "Two metrics with different scales on the same categories often need mixed encodings plus a second axis.",
        "When one series should read as volume and another as rate, combining geometries avoids flattening the smaller signal.",
        "Dual-axis composites keep one canvas instead of forcing readers to mentally stitch two separate charts.",
    ),
}

def _diverse_chart_reason(
    description: str,
    chart_type: str,
    rank_index: int,
    confidence: float,
    matched_keywords: List[str],
    language: str,
) -> str:
    ct = (chart_type or "bar").lower()
    pool = (_REASON_EN if language == "en" else SZH.REASON_POOL_ZH).get(ct) or (
        _REASON_EN["bar"] if language == "en" else SZH.REASON_POOL_ZH["bar"]
    )
    seed = hashlib.md5(
        f"{rank_index}|{ct}|{(description or '')[:420]}".encode("utf-8", errors="ignore")
    ).hexdigest()
    idx = int(seed, 16) % len(pool)
    body = pool[idx]
    openings = _RUNNER_OPENING_EN if language == "en" else SZH.RUNNER_OPENING_ZH
    if rank_index > 0:
        oi = int(seed[8:16], 16) % len(openings)
        body = f"{openings[oi]}{body}"
    kws = [k for k in (matched_keywords or []) if k]
    if kws:
        clip = ", ".join(kws[:2])
        if language == "en":
            body = f"{body} (Hooks: {clip}.)"
        else:
            body = f"{body}{SZH.fmt_reason_body_hook_zh(clip)}"
    return body


def _parse_reasons_array_from_llm(content: str) -> Optional[List[str]]:
    raw = (content or "").strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```[a-zA-Z0-9_-]*\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw).strip()
    j0 = raw.find("{")
    j1 = raw.rfind("}") + 1
    if j0 < 0 or j1 <= j0:
        return None
    try:
        obj = json.loads(raw[j0:j1])
    except json.JSONDecodeError:
        return None
    arr = obj.get("reasons")
    if not isinstance(arr, list):
        return None
    out: List[str] = []
    for x in arr:
        if isinstance(x, str) and x.strip():
            out.append(x.strip())
    return out or None


def enrich_chart_recommendation_reasons_llm(
    description: str,
    recommendations: List[Dict[str, Any]],
    language: str,
) -> Optional[List[str]]:
    """One batched LLM call: distinct prose per row; must not change ordering."""
    if not recommendations:
        return None
    key = (DEEPSEEK_API_KEY or "").strip()
    if not key:
        return None

    rows: List[Dict[str, Any]] = []
    for i, rec in enumerate(recommendations):
        rows.append(
            {
                "rank": i + 1,
                "type": rec.get("type"),
                "chart_name": rec.get("name"),
                "keywords": rec.get("matched_keywords") or [],
            }
        )
    desc_short = (description or "")[:1400]
    n = len(recommendations)
    if language == "en":
        user = f"""Data excerpt (may contain irrelevant prose):
---
{desc_short}
---

The chart-type ranking below is FINAL from the server (rank 1 is already used to render the chart). For EACH row, write exactly ONE sentence (≤40 words) that explains why that chart idiom is a sensible reading of the excerpt. Vary openings and rhythm across rows; do not reuse the same first three words. Do not contradict the order, do not suggest a different winner, do not mention confidence percentages.

Rows:
{json.dumps(rows, ensure_ascii=False)}

Return JSON only. The "reasons" array MUST contain exactly {n} strings, e.g. {{"reasons": ["…", "…"]}}."""
        system = (
            "You write crisp chart rationales for a product UI. Output valid JSON only—no markdown fences, no commentary."
        )
    else:
        user = PZH.enrich_recommendation_reasons_user_zh(desc_short, rows, n)
        system = PZH.enrich_recommendation_reasons_system_zh()

    payload = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": 0.72,
        "max_tokens": min(1200, 180 * n + 120),
    }
    headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}
    try:
        resp = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=22)
        resp.raise_for_status()
        content = resp.json()["choices"][0]["message"]["content"]
    except Exception as exc:
        print(f"enrich_chart_recommendation_reasons_llm failed: {exc}")
        return None

    parsed = _parse_reasons_array_from_llm(content)
    if not parsed or len(parsed) != n:
        return None
    return parsed


def build_generation_aligned_analysis_summary(language: str) -> Dict[str, str]:
    """Short copy for the UI when recommendations are tied to the generation scorer."""
    if language == "en":
        return {
            "data_characteristics": "We skim your excerpt for numeric evidence and semantic cues—noise is expected.",
            "visualization_goals": "The top card mirrors what was actually rendered; the lines below are alternate lenses on the same text.",
            "recommended_approach": "Tap Apply type to re-run with another encoding; each rationale is rewritten to feel human, while the ranking stays tied to the server scorer.",
        }
    return dict(SZH.ANALYSIS_SUMMARY_UI_ZH)


def build_chart_type_recommendations_from_scores(
    type_scores: Dict[str, float],
    selected_chart_type: str,
    description: str,
    language: str,
    top_n: int = 5,
) -> List[Dict[str, Any]]:
    """
    Build ranked recommendations from the same type_scores used to pick chart_type,
    so UI #1 always matches the rendered chart.
    """
    scores: Dict[str, float] = {}
    for k, v in (type_scores or {}).items():
        if isinstance(v, (int, float)):
            scores[str(k)] = float(v)
    if not scores:
        return []

    ranked = sorted(scores.items(), key=lambda kv: (kv[1], kv[0]), reverse=True)
    sel = normalize_chart_type_code(selected_chart_type or "") or ""
    if not sel or sel not in scores:
        sel = ranked[0][0]

    rest = [(t, s) for t, s in ranked if t != sel]
    ordered = [(sel, max(0.0, scores.get(sel, 0.0)))] + rest[: max(0, top_n - 1)]
    ordered = ordered[:top_n]

    total = sum(s for _, s in ordered)
    if total <= 0:
        mkw = _infer_matched_keywords(description, sel, language=language)
        one = {
            "type": sel,
            "name": get_chart_type_name(sel, language=language),
            "confidence": 100.0,
            "reason": _diverse_chart_reason(description, sel, 0, 100.0, mkw, language),
            "matched_keywords": mkw,
        }
        alt = enrich_chart_recommendation_reasons_llm(description, [one], language)
        if alt and len(alt) == 1:
            one["reason"] = alt[0]
        return [one]

    raw_pcts = [(t, (max(0.0, s) / total) * 100.0) for t, s in ordered]
    rounded: List[Tuple[str, float]] = []
    for t, p in raw_pcts:
        rounded.append((t, round(p, 1)))
    drift = round(100.0 - sum(p for _, p in rounded), 1)
    if rounded and abs(drift) >= 0.05:
        t0, p0 = rounded[0]
        rounded[0] = (t0, round(p0 + drift, 1))

    out: List[Dict[str, Any]] = []
    for rank_index, (chart_type, confidence) in enumerate(rounded):
        mkw = _infer_matched_keywords(description, chart_type, language=language)
        out.append(
            {
                "type": chart_type,
                "name": get_chart_type_name(chart_type, language=language),
                "confidence": confidence,
                "reason": _diverse_chart_reason(
                    description, chart_type, rank_index, confidence, mkw, language
                ),
                "matched_keywords": mkw,
            }
        )
    llm_reasons = enrich_chart_recommendation_reasons_llm(description, out, language)
    if llm_reasons and len(llm_reasons) == len(out):
        for i, rr in enumerate(llm_reasons):
            out[i]["reason"] = rr
    return out
